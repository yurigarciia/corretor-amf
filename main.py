"""
CorretorAMF — Backend FastAPI
Verifica conformidade de trabalhos acadêmicos com as normas da AMF via Google Gemini.
"""

import io
import json
import logging
import os
import re
from collections import Counter
from contextlib import asynccontextmanager
from pathlib import Path

import fitz  # pymupdf
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

from prompts import SYSTEM_PROMPT, build_user_prompt

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
MODEL_NAME = "gemini-3-flash-preview"
MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
ALLOWED_MIME_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
}


@asynccontextmanager
async def lifespan(app: FastAPI):
    if not GEMINI_API_KEY:
        logger.warning("GEMINI_API_KEY not set — /analyze will fail at runtime")
    else:
        genai.configure(api_key=GEMINI_API_KEY)
        logger.info("Gemini SDK configured with model: %s", MODEL_NAME)
    yield


app = FastAPI(
    title="CorretorAMF",
    description="Verificador de conformidade de trabalhos acadêmicos com normas AMF via IA",
    version="1.0.0",
    lifespan=lifespan,
)

static_dir = Path(__file__).parent / "static"
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")


# ---------------------------------------------------------------------------
# Helpers — DOCX
# ---------------------------------------------------------------------------

def _resolve_font_name(doc, style) -> str | None:
    """Percorre a hierarquia de estilos para encontrar o primeiro nome de fonte explícito."""
    seen: set = set()
    while style and style.name not in seen:
        seen.add(style.name)
        fname = getattr(style.font, "name", None)
        # Ignora referências de tema como "+Body" ou "+mn-lt"
        if fname and not fname.startswith("+"):
            return fname
        style = getattr(style, "base_style", None)
    return None


def _describe_line_spacing(pf) -> str | None:
    """Retorna descrição legível do espaçamento entre linhas de um parágrafo."""
    try:
        rule = pf.line_spacing_rule
        ls = pf.line_spacing
        if rule == WD_LINE_SPACING.SINGLE:
            return "simples (1,0)"
        if rule == WD_LINE_SPACING.ONE_POINT_FIVE:
            return "1,5 linhas"
        if rule == WD_LINE_SPACING.DOUBLE:
            return "duplo (2,0)"
        if rule == WD_LINE_SPACING.EXACTLY and ls:
            return f"exato {ls.pt:.1f}pt"
        if rule == WD_LINE_SPACING.MULTIPLE and ls:
            return f"múltiplo {float(ls):.2f}x"
        if rule == WD_LINE_SPACING.AT_LEAST and ls:
            return f"mínimo {ls.pt:.1f}pt"
    except Exception:
        pass
    return None


def _extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    out: list[str] = []

    out.append("╔══════════════════════════════════════╗")
    out.append("║  METADADOS DE FORMATAÇÃO — DOCX      ║")
    out.append("╚══════════════════════════════════════╝")

    # ── Margens e tamanho de página ───────────────────────────
    for i, sec in enumerate(doc.sections, 1):
        label = f"Seção {i}" if len(doc.sections) > 1 else "Documento"

        def _cm(v):
            try:
                return f"{v.cm:.2f} cm"
            except Exception:
                return "não definido"

        out.append(f"\nMARGENS ({label}):")
        out.append(f"  Superior : {_cm(sec.top_margin)}")
        out.append(f"  Inferior : {_cm(sec.bottom_margin)}")
        out.append(f"  Esquerda : {_cm(sec.left_margin)}")
        out.append(f"  Direita  : {_cm(sec.right_margin)}")
        if sec.page_width and sec.page_height:
            out.append(f"  Tamanho da página: {sec.page_width.cm:.1f} × {sec.page_height.cm:.1f} cm")
        if sec.header_distance:
            out.append(f"  Distância cabeçalho: {_cm(sec.header_distance)}")
        if sec.footer_distance:
            out.append(f"  Distância rodapé: {_cm(sec.footer_distance)}")

    # ── Análise tipográfica ───────────────────────────────────
    font_chars: Counter = Counter()
    size_chars: Counter = Counter()
    spacing_counter: Counter = Counter()
    first_line_indents: list[float] = []
    heading_samples: dict = {}

    # Espaçamento do estilo Normal (herança base)
    try:
        normal_style = doc.styles["Normal"]
        normal_ls = _describe_line_spacing(normal_style.paragraph_format)
        if normal_ls:
            out.append(f"\nEstilo Normal (padrão do documento): espaçamento {normal_ls}")
        # Fonte padrão do Normal
        normal_font = _resolve_font_name(doc, normal_style)
        if normal_font:
            out.append(f"Fonte padrão (estilo Normal): {normal_font}")
        if normal_style.font.size:
            try:
                out.append(f"Tamanho padrão (estilo Normal): {normal_style.font.size.pt:.0f}pt")
            except Exception:
                pass
    except Exception:
        pass

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        pf = para.paragraph_format
        style = para.style

        # Espaçamento: tenta no parágrafo, depois no estilo
        ls = _describe_line_spacing(pf)
        if not ls and style:
            ls = _describe_line_spacing(style.paragraph_format)
        if ls:
            spacing_counter[ls] += len(text)

        # Recuo de primeira linha
        fi = pf.first_line_indent
        if fi:
            try:
                val = fi.cm
                if val > 0.05:
                    first_line_indents.append(round(val, 2))
            except Exception:
                pass

        # Detecção de nível de heading pelo nome do estilo
        sname = style.name if style else ""
        heading_level = None
        for lvl in range(1, 7):
            if f"Heading {lvl}" in sname:
                heading_level = lvl
                break

        for run in para.runs:
            if not run.text.strip():
                continue

            # Nome da fonte: run → estilo → hierarquia de estilos
            fname = run.font.name
            if not fname or fname.startswith("+"):
                fname = _resolve_font_name(doc, style)
            if fname:
                font_chars[fname] += len(run.text)

            # Tamanho: run → estilo
            fsize = run.font.size
            if not fsize and style:
                fsize = getattr(style.font, "size", None)
            if fsize:
                try:
                    pt = round(fsize.pt)
                    size_chars[pt] += len(run.text)
                    if heading_level is not None and heading_level not in heading_samples:
                        bold = run.bold
                        if bold is None and style:
                            bold = getattr(style.font, "bold", None)
                        heading_samples[heading_level] = (fname or "N/D", pt, bool(bold))
                except Exception:
                    pass

    out.append("\nTIPOGRAFIA (extraída dos metadados do arquivo):")
    if font_chars:
        out.append("  Fontes detectadas (por volume de caracteres):")
        for fname, cnt in font_chars.most_common(6):
            out.append(f"    • {fname}: {cnt} chars")
    else:
        out.append("  Fontes: não detectadas diretamente nos runs (possível herança de tema)")

    if size_chars:
        out.append("  Tamanhos de fonte detectados:")
        for sz, cnt in size_chars.most_common(8):
            out.append(f"    • {sz}pt: {cnt} chars")

    out.append("\nESPAÇAMENTO ENTRE LINHAS (por volume de texto):")
    if spacing_counter:
        for sp, cnt in spacing_counter.most_common(5):
            out.append(f"  • {sp}: {cnt} chars")
    else:
        out.append("  Não detectado explicitamente nos parágrafos (pode estar herdado do tema/Normal)")

    if first_line_indents:
        avg = sum(first_line_indents) / len(first_line_indents)
        out.append(f"\nRecuo de primeira linha (média amostral): {avg:.2f} cm")

    if heading_samples:
        out.append("\nEstilos de títulos detectados:")
        for lvl in sorted(heading_samples):
            fname, pt, bold = heading_samples[lvl]
            out.append(f"  • Heading {lvl}: {fname} {pt}pt {'negrito' if bold else 'normal'}")

    # ── Conteúdo com marcadores de estrutura ─────────────────
    out.append("\n╔══════════════════════════════════════╗")
    out.append("║       CONTEÚDO DO DOCUMENTO          ║")
    out.append("╚══════════════════════════════════════╝")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        sname = para.style.name if para.style else ""
        prefix = ""
        for lvl in range(1, 7):
            if f"Heading {lvl}" in sname:
                prefix = f"[H{lvl}] "
                break
        if not prefix:
            if "Title" in sname:
                prefix = "[TÍTULO] "
            elif "Subtitle" in sname:
                prefix = "[SUBTÍTULO] "
            else:
                runs_with_text = [r for r in para.runs if r.text.strip()]
                if runs_with_text and all(r.bold for r in runs_with_text):
                    prefix = "[NEGRITO] "

        out.append(f"{prefix}{text}")

    for table in doc.tables:
        out.append("\n[TABELA]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                out.append("  " + " | ".join(cells))

    return "\n".join(out)


# ---------------------------------------------------------------------------
# Helpers — PDF
# ---------------------------------------------------------------------------

def _clean_font_name(raw: str) -> str:
    """Remove prefixo de subset (ABCDEF+) e sufixos comuns de nomes de fonte PDF."""
    name = raw.split("+")[-1] if "+" in raw else raw
    for sfx in ["-Bold", "-Italic", "-BoldItalic", "-Regular", "MT", "PS", "-Roman", "-Light", "-Medium"]:
        name = name.replace(sfx, "")
    return name.strip()


def _extract_pdf_text(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    out: list[str] = []

    font_chars: Counter = Counter()
    size_chars: Counter = Counter()
    page_margins: list[dict] = []
    # Armazena conteúdo por página: lista de (texto, tamanho, is_bold)
    content_pages: list[tuple[int, list[tuple[str, float, bool]]]] = []

    first_page_size: tuple[float, float] | None = None

    for page_num in range(doc.page_count):
        page = doc[page_num]
        pw = page.rect.width
        ph = page.rect.height

        if page_num == 0:
            first_page_size = (pw * 0.0352778, ph * 0.0352778)

        page_dict = page.get_text(
            "dict",
            flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE,
        )

        min_x, max_x = pw, 0.0
        min_y, max_y = ph, 0.0
        page_lines: list[tuple[str, float, bool]] = []

        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:  # 0 = bloco de texto
                continue
            bx0, by0, bx1, by1 = block["bbox"]
            # Ignora fragmentos minúsculos (artefatos)
            if (bx1 - bx0) < 5 or (by1 - by0) < 2:
                continue
            min_x = min(min_x, bx0)
            max_x = max(max_x, bx1)
            min_y = min(min_y, by0)
            max_y = max(max_y, by1)

            for line in block.get("lines", []):
                line_parts: list[str] = []
                line_max_size = 0.0
                line_has_bold = False

                for span in line.get("spans", []):
                    stext = span.get("text", "")
                    ssize = span.get("size", 0.0)
                    sflags = span.get("flags", 0)
                    sfont = span.get("font", "")

                    if not stext.strip():
                        continue

                    line_parts.append(stext)
                    line_max_size = max(line_max_size, ssize)
                    if sflags & 16:  # bit 4 = bold
                        line_has_bold = True

                    # Contadores de fonte e tamanho
                    clean = _clean_font_name(sfont)
                    if clean:
                        font_chars[clean] += len(stext.strip())
                    if ssize > 0:
                        size_chars[round(ssize)] += len(stext.strip())

                full_line = "".join(line_parts).strip()
                if full_line:
                    page_lines.append((full_line, line_max_size, line_has_bold))

        # Margens inferidas (primeiras 5 páginas)
        if page_num < 5 and min_x < pw and min_y < ph:
            page_margins.append({
                "page": page_num + 1,
                "top": round(min_y * 0.0352778, 2),
                "bottom": round((ph - max_y) * 0.0352778, 2),
                "left": round(min_x * 0.0352778, 2),
                "right": round((pw - max_x) * 0.0352778, 2),
            })

        content_pages.append((page_num + 1, page_lines))

    doc.close()

    # ── Cabeçalho de metadados ────────────────────────────────
    out.append("╔══════════════════════════════════════╗")
    out.append("║   METADADOS DE FORMATAÇÃO — PDF      ║")
    out.append("╚══════════════════════════════════════╝")
    out.append(f"Total de páginas: {len(content_pages)}")

    if first_page_size:
        w, h = first_page_size
        out.append(f"Tamanho da página: {w:.1f} × {h:.1f} cm  (A4 padrão = 21,0 × 29,7 cm)")

    if page_margins:
        out.append("\nMARGENS INFERIDAS (posição dos blocos de texto vs. borda da página):")
        out.append("  ATENÇÃO: valores aproximados com imprecisão de ±0,5 cm — o texto não")
        out.append("  preenche até a borda exata da margem. Use tolerância de ±0,5 cm ao avaliar.")
        out.append("  Cabeçalhos e rodapés podem reduzir artificialmente sup/inf.")
        for pm in page_margins:
            out.append(
                f"  Pág {pm['page']}: "
                f"sup {pm['top']:.2f} cm | "
                f"inf {pm['bottom']:.2f} cm | "
                f"esq {pm['left']:.2f} cm | "
                f"dir {pm['right']:.2f} cm"
            )

    out.append("\nTIPOGRAFIA (extraída dos metadados do PDF):")
    if font_chars:
        out.append("  Fontes detectadas (por volume de caracteres):")
        for fname, cnt in font_chars.most_common(6):
            out.append(f"    • {fname}: {cnt} chars")
    if size_chars:
        out.append("  Tamanhos de fonte detectados:")
        for sz, cnt in size_chars.most_common(8):
            out.append(f"    • {sz}pt: {cnt} chars")

    # ── Conteúdo com marcadores estruturais ──────────────────
    out.append("\n╔══════════════════════════════════════╗")
    out.append("║       CONTEÚDO DO DOCUMENTO          ║")
    out.append("╚══════════════════════════════════════╝")

    body_size = size_chars.most_common(1)[0][0] if size_chars else 12

    for page_num, lines in content_pages:
        out.append(f"\n--- Página {page_num} ---")
        for line_text, line_size, line_bold in lines:
            prefix = ""
            if line_size > body_size + 3:
                prefix = f"[TÍTULO {round(line_size)}pt] "
            elif line_bold and line_size >= body_size:
                prefix = "[NEGRITO] "
            out.append(f"{prefix}{line_text}")

    return "\n".join(out)


# ---------------------------------------------------------------------------
# Helpers — Gemini
# ---------------------------------------------------------------------------

def _parse_gemini_response(raw: str) -> dict:
    """
    Extrai JSON da resposta do Gemini.
    O modelo às vezes envolve o JSON em markdown mesmo sendo instruído a não fazer isso.
    """
    text = raw.strip()

    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
        text = text.strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass

    # Fallback: tenta reparar JSON truncado (ex: MAX_TOKENS cortou no meio)
    repaired = _repair_truncated_json(text)
    if repaired:
        logger.warning("JSON reparado após truncamento da resposta do modelo")
        return repaired

    raise ValueError(f"Não foi possível extrair JSON válido da resposta do modelo: {raw[:200]}")


def _repair_truncated_json(text: str) -> dict | None:
    """
    Tenta recuperar um JSON truncado (ex: cortado por MAX_TOKENS) fechando
    as estruturas abertas a partir do último objeto completo encontrado.
    """
    # Remove markdown residual
    text = re.sub(r"^```(?:json)?\s*", "", text.strip())
    text = re.sub(r"\s*```$", "", text).strip()

    last_brace = text.rfind("}")
    if last_brace == -1:
        return None

    base = text[: last_brace + 1]
    for closing in [
        ', "suggestions": []}',
        ', "suggestions": [], "score": "0%"}',
        "]}",
    ]:
        try:
            result = json.loads(base + closing)
            result.setdefault("score", "0%")
            result.setdefault("criteria", [])
            result.setdefault("suggestions", [])
            return result
        except json.JSONDecodeError:
            pass

    return None


def _call_gemini_with_text(document_text: str, extra_context: str) -> dict:
    """Chama o Gemini com o conteúdo extraído do documento como texto."""
    MAX_INPUT_CHARS = 80_000
    if len(document_text) > MAX_INPUT_CHARS:
        document_text = document_text[:MAX_INPUT_CHARS] + "\n\n[... documento truncado para análise ...]"

    model = genai.GenerativeModel(
        model_name=MODEL_NAME,
        system_instruction=SYSTEM_PROMPT,
    )
    user_prompt = build_user_prompt(document_text, extra_context)
    response = model.generate_content(
        user_prompt,
        generation_config=genai.GenerationConfig(
            temperature=0.1,
            max_output_tokens=16384,
        ),
        request_options={"timeout": 120},
    )

    candidate = response.candidates[0] if response.candidates else None
    if candidate and candidate.finish_reason.name != "STOP":
        logger.warning("Gemini finish_reason: %s", candidate.finish_reason.name)

    return _parse_gemini_response(response.text)


# ---------------------------------------------------------------------------
# Rotas
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    """Serve o frontend principal."""
    index_path = static_dir / "index.html"
    return HTMLResponse(content=index_path.read_text(encoding="utf-8"))


@app.get("/health")
async def health_check():
    """Verifica se o serviço está operacional."""
    return {"status": "ok", "model": MODEL_NAME}


@app.post("/analyze")
async def analyze_document(
    file: UploadFile = File(..., description="PDF ou DOCX do trabalho acadêmico"),
    context: str = Form(default="", description="Contexto adicional opcional informado pelo aluno"),
):
    """
    Analisa um trabalho acadêmico e retorna um relatório de conformidade com as normas AMF.

    - **file**: arquivo PDF ou DOCX, máximo 10 MB
    - **context**: informações extras sobre o trabalho (tipo, disciplina, restrições aplicáveis)
    """
    if not GEMINI_API_KEY:
        raise HTTPException(
            status_code=500,
            detail="Serviço não configurado: GEMINI_API_KEY ausente. Contate o administrador.",
        )

    content_type = file.content_type or ""
    content_type_base = content_type.split(";")[0].strip()
    filename = file.filename or ""

    if content_type_base not in ALLOWED_MIME_TYPES:
        if filename.lower().endswith(".pdf"):
            content_type_base = "application/pdf"
        elif filename.lower().endswith(".docx"):
            content_type_base = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Tipo de arquivo não suportado: '{content_type}'. Envie um PDF ou DOCX.",
            )

    file_bytes = await file.read()

    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="Arquivo enviado está vazio.")

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        size_mb = len(file_bytes) / (1024 * 1024)
        raise HTTPException(
            status_code=413,
            detail=f"Arquivo muito grande: {size_mb:.1f} MB. Limite: {MAX_FILE_SIZE_MB} MB.",
        )

    logger.info(
        "Analisando arquivo: %s | tipo: %s | tamanho: %.1f KB",
        filename,
        content_type_base,
        len(file_bytes) / 1024,
    )

    try:
        if content_type_base == "application/pdf":
            document_text = _extract_pdf_text(file_bytes)
        else:
            document_text = _extract_docx_text(file_bytes)

        if not document_text.strip():
            raise HTTPException(
                status_code=422,
                detail="Não foi possível extrair conteúdo do arquivo. O arquivo pode estar corrompido ou protegido.",
            )

        result = _call_gemini_with_text(document_text, context)

    except HTTPException:
        raise
    except Exception as e:
        err_str = str(e)
        if "429" in err_str or "quota" in err_str.lower() or "resource_exhausted" in err_str.lower():
            raise HTTPException(
                status_code=429,
                detail="Limite de requisições da IA atingido. O sistema usa uma cota diária gratuita — aguarde alguns minutos e tente novamente.",
            )
        logger.exception("Erro ao processar documento")
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao analisar o documento: {err_str}. Tente novamente.",
        )

    result.setdefault("score", "0%")
    result.setdefault("criteria", [])
    result.setdefault("suggestions", [])

    return JSONResponse(content=result)
